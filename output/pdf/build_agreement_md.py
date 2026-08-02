from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

source = Path('/Users/anurag/Desktop/Agreement.md')
outdir = Path('/Users/anurag/Desktop/Logistics Vision AI/output/pdf')
outdir.mkdir(parents=True, exist_ok=True)
docx_path = outdir / 'Agreement_Client_Ready.docx'

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Sama Devanagari'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Sama Devanagari')
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.1

for name, size in [('Title', 19), ('Heading 1', 15), ('Heading 2', 12.5), ('Heading 3', 11.5)]:
    st = styles[name]
    st.font.name = 'Sama Devanagari'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Sama Devanagari')
    st.font.size = Pt(size)
    st.font.bold = True

def font_run(run, size=None, bold=None):
    run.font.name = 'Sama Devanagari'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Sama Devanagari')
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run('Agreement | पृष्ठ ')
    font_run(r, 8)
    fld = r._r.get_or_add_rPr()
    # PAGE field is kept as a regular field sequence for LibreOffice compatibility.
    from docx.oxml import OxmlElement
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r._r.append(begin); r._r.append(instr); r._r.append(end)

for raw in source.read_text(encoding='utf-8').splitlines():
    line = raw.rstrip()
    if not line.strip():
        doc.add_paragraph('')
        continue
    s = line.strip()
    if s.startswith('### '):
        p = doc.add_paragraph(style='Heading 2')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s[4:]); font_run(r, 12.5, True)
    elif s.startswith('## '):
        p = doc.add_paragraph(style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s[3:]); font_run(r, 15, True)
    elif s.startswith('# '):
        p = doc.add_paragraph(style='Title')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s[2:]); font_run(r, 19, True)
    elif s.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(s[2:]); font_run(r)
    elif s[:1].isdigit() and '. ' in s[:5]:
        # Preserve the source numbering exactly; do not let Word continue a
        # single automatic list across separate clauses and annexures.
        p = doc.add_paragraph()
        r = p.add_run(s); font_run(r)
    else:
        p = doc.add_paragraph()
        r = p.add_run(s); font_run(r)

doc.save(docx_path)
print(docx_path)
