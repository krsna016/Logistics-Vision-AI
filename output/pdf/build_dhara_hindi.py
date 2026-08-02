from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

base = Path(__file__).parent
src = base / 'dhara_agreement_hindi.txt'
docx_path = base / 'DHARA_AGREEMENT_HINDI.docx'

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.68)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.78)
sec.right_margin = Inches(0.78)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Sama Devanagari'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Sama Devanagari')
normal.font.size = Pt(10.2)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.08

for name in ['Title', 'Heading 1', 'Heading 2']:
    st = styles[name]
    st.font.name = 'Sama Devanagari'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Sama Devanagari')
    st.font.bold = True

styles['Title'].font.size = Pt(18)
styles['Heading 1'].font.size = Pt(13)
styles['Heading 2'].font.size = Pt(11)

def set_run_font(run, size=None, bold=None):
    run.font.name = 'Sama Devanagari'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Sama Devanagari')
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold

def add_page_number(paragraph):
    run = paragraph.add_run('पृष्ठ ')
    set_run_font(run, 8)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

lines = src.read_text(encoding='utf-8').splitlines()
for line in lines:
    s = line.strip()
    if not s:
        doc.add_paragraph('')
        continue
    p = doc.add_paragraph()
    # Keep contract clauses and annexure headings visibly distinct while retaining all text.
    is_annex = s.startswith('ANNEXURE')
    is_heading = (len(s) < 75 and (s[:2].isdigit() and '.' in s[:4] or s.isupper()))
    if is_annex:
        p.style = 'Heading 1'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif s in ('समझौता', 'भारत गैर-न्यायिक', 'ई-स्टाम्प प्रमाणपत्र'):
        p.style = 'Title' if s == 'समझौता' else 'Heading 2'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif is_heading:
        p.style = 'Heading 2'
    run = p.add_run(s)
    set_run_font(run, 12 if is_annex else (13 if s == 'समझौता' else 10.2), is_annex or s == 'समझौता')

doc.save(docx_path)
print(docx_path)
